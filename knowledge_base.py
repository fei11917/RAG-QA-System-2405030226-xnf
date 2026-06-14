import os
import re
from typing import List

class SimpleTextSplitter:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end < len(text):
                gap = text.rfind(' ', start, end)
                if gap != -1 and gap > start + self.chunk_size // 2:
                    end = gap
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks

class KnowledgeBase:
    def __init__(self, 
                 persist_directory: str = "./chroma_db",
                 embedding_model: str = "nomic-embed-text"):
        self.persist_directory = persist_directory
        self.embedding_model = embedding_model
        self.vectorstore = None
        self.text_splitter = SimpleTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
        )
        self._embeddings = None
        
    @property
    def embeddings(self):
        if self._embeddings is None:
            from langchain_community.embeddings import OllamaEmbeddings
            self._embeddings = OllamaEmbeddings(model=self.embedding_model)
        return self._embeddings
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Failed to read PDF {file_path}: {str(e)}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            from docx import Document
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Failed to read DOCX {file_path}: {str(e)}")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            print(f"Unsupported format: {ext}")
            return ""
    
    def load_documents_from_folder(self, folder_path: str) -> list:
        documents = []
        
        if not os.path.exists(folder_path):
            print(f"Folder not found: {folder_path}")
            return documents
        
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            if os.path.isfile(file_path):
                ext = os.path.splitext(filename)[1].lower()
                
                if ext in ['.pdf', '.docx']:
                    print(f"Processing file: {filename}")
                    text = self.extract_text_from_file(file_path)
                    
                    if text.strip():
                        from langchain_core.documents import Document as LangDocument
                        doc = LangDocument(
                            page_content=text,
                            metadata={"source": filename, "file_path": file_path}
                        )
                        documents.append(doc)
        
        return documents
    
    def load_documents_from_uploaded_files(self, uploaded_files) -> list:
        documents = []
        
        for uploaded_file in uploaded_files:
            filename = uploaded_file.name
            ext = os.path.splitext(filename)[1].lower()
            
            if ext in ['.pdf', '.docx']:
                print(f"Processing uploaded file: {filename}")
                
                temp_path = os.path.join("./temp_docs", filename)
                os.makedirs("./temp_docs", exist_ok=True)
                
                with open(temp_path, 'wb') as f:
                    f.write(uploaded_file.getbuffer())
                
                text = self.extract_text_from_file(temp_path)
                
                if text.strip():
                    from langchain_core.documents import Document as LangDocument
                    doc = LangDocument(
                        page_content=text,
                        metadata={"source": filename, "file_path": temp_path}
                    )
                    documents.append(doc)
        
        return documents
    
    def split_documents(self, documents: list) -> list:
        if not documents:
            return []
        
        split_docs = []
        for doc in documents:
            chunks = self.text_splitter.split_text(doc.page_content)
            for chunk in chunks:
                from langchain_core.documents import Document as LangDocument
                split_docs.append(LangDocument(
                    page_content=chunk,
                    metadata=doc.metadata
                ))
        
        print(f"Split completed: {len(documents)} docs -> {len(split_docs)} chunks")
        return split_docs
    
    def build_vectorstore(self, documents: list):
        if not documents:
            print("No documents to add")
            return
        
        split_docs = self.split_documents(documents)
        
        if self.vectorstore is None:
            from langchain_community.vectorstores import Chroma
            self.vectorstore = Chroma.from_documents(
                documents=split_docs,
                embedding=self.embeddings,
                persist_directory=self.persist_directory
            )
        else:
            self.vectorstore.add_documents(split_docs)
        
        print(f"Vectorstore updated, total {len(split_docs)} chunks")
    
    def load_vectorstore(self):
        try:
            from langchain_community.vectorstores import Chroma
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
            print("Vectorstore loaded successfully")
            return True
        except Exception as e:
            print(f"Failed to load vectorstore: {str(e)}")
            return False
    
    def search(self, query: str, k: int = 3) -> list:
        if self.vectorstore is None:
            print("Vectorstore not initialized")
            return []
        
        results = self.vectorstore.similarity_search(query, k=k)
        return results
    
    def get_retriever(self, k: int = 3):
        if self.vectorstore is None:
            return None
        return self.vectorstore.as_retriever(search_kwargs={"k": k})
    
    def get_document_count(self) -> int:
        if self.vectorstore is None:
            return 0
        
        try:
            collection = self.vectorstore._collection
            return collection.count()
        except:
            return 0
    
    def clear_vectorstore(self):
        import shutil
        if os.path.exists(self.persist_directory):
            shutil.rmtree(self.persist_directory)
            print("Vectorstore cleared")
        self.vectorstore = None


if __name__ == "__main__":
    kb = KnowledgeBase()
    
    docs_folder = "./documents"
    
    if os.path.exists(docs_folder):
        documents = kb.load_documents_from_folder(docs_folder)
        
        if documents:
            kb.build_vectorstore(documents)
            
            test_query = "What is natural language processing"
            results = kb.search(test_query, k=3)
            
            print("\n" + "=" * 50)
            print(f"Test query: {test_query}")
            print("=" * 50)
            
            for i, doc in enumerate(results, 1):
                print(f"\nResult {i}:")
                print(f"Source: {doc.metadata.get('source', 'Unknown')}")
                print(f"Content: {doc.page_content[:200]}...")
    else:
        print(f"Please create {docs_folder} folder and add PDF or DOCX documents")
